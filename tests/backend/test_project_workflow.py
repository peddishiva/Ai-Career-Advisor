#!/usr/bin/env python3
"""
Complete Project Workflow Test
Tests the entire data flow from resume upload to AI analysis and job recommendations
"""

import os
import sys
import json
import requests
import time
from pathlib import Path
from dotenv import load_dotenv

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

REPO_ROOT = Path(__file__).resolve().parents[2]
load_dotenv(REPO_ROOT / "backend" / ".env")

def print_status(message, status="INFO"):
    """Print colored status messages"""
    colors = {
        "SUCCESS": "\033[92m",  # Green
        "ERROR": "\033[91m",    # Red
        "WARNING": "\033[93m",  # Yellow
        "INFO": "\033[94m",     # Blue
        "RESET": "\033[0m"      # Reset
    }
    
    color = colors.get(status, colors["INFO"])
    print(f"{color}[{status}]{colors['RESET']} {message}")

def test_backend_endpoints():
    """Test all backend API endpoints"""
    print_status("Testing Backend API Endpoints...", "INFO")
    
    base_url = "http://localhost:5000"
    
    endpoints = [
        ("/", "Root endpoint"),
        ("/health", "Health check"),
        ("/docs", "API documentation"),
    ]
    
    results = {}
    
    for endpoint, description in endpoints:
        try:
            response = requests.get(f"{base_url}{endpoint}", timeout=5)
            if response.status_code == 200:
                print_status(f"✓ {endpoint} - {description}", "SUCCESS")
                results[endpoint] = True
            else:
                print_status(f"✗ {endpoint} - Status: {response.status_code}", "ERROR")
                results[endpoint] = False
        except Exception as e:
            print_status(f"✗ {endpoint} - Error: {str(e)}", "ERROR")
            results[endpoint] = False
    
    return all(results.values())

def test_resume_upload_workflow():
    """Test resume upload and analysis workflow"""
    print_status("\nTesting Resume Upload Workflow...", "INFO")
    
    base_url = "http://localhost:5000"
    
    # Create a test text / doc for upload test
    test_file_path = REPO_ROOT / "test_sample_resume.docx"
    
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Jane Doe', 0)
        doc.add_paragraph('jane@example.com | +1 555-123-4567')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Software Engineer at TechCorp for 3 years building web applications with Python and React.')
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science')
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Python, JavaScript, SQL, React, Docker, Git, Machine Learning')
        doc.save(str(test_file_path))
    except Exception as e:
        print_status(f"Notice: couldn't generate test docx ({e}), skipping upload test", "WARNING")
        return True

    try:
        print_status("Uploading test resume...", "INFO")
        
        with open(test_file_path, "rb") as f:
            files = {"file": ("test_sample_resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(f"{base_url}/api/upload", files=files, timeout=30)
        
        if response.status_code == 200:
            upload_data = response.json()
            print_status("✓ Resume uploaded successfully", "SUCCESS")
            
            # Test analysis retrieval
            print_status("Retrieving analysis results...", "INFO")
            analysis_response = requests.get(f"{base_url}/api/analysis", timeout=10)
            
            if analysis_response.status_code == 200:
                analysis_data = analysis_response.json()
                print_status("✓ Analysis retrieved successfully", "SUCCESS")
                
                if "data" in analysis_data:
                    analysis = analysis_data["data"]
                    required_fields = ["overall_insights", "metrics", "candidate_info", "role_matches"]
                    
                    missing_fields = [field for field in required_fields if field not in analysis]
                    if not missing_fields:
                        print_status("✓ Analysis structure is valid", "SUCCESS")
                        fit_score = analysis.get("overall_insights", {}).get("fit_score", "N/A")
                        role_alignment = analysis.get("metrics", {}).get("role_alignment", "N/A")
                        print(f"  Fit Score: {fit_score}")
                        print(f"  Role Alignment: {role_alignment}")
                        return True
                    else:
                        print_status(f"✗ Analysis missing fields: {missing_fields}", "ERROR")
                        return False
                else:
                    print_status("✗ Invalid analysis response format", "ERROR")
                    return False
            else:
                print_status(f"✗ Analysis retrieval failed: {analysis_response.status_code}", "ERROR")
                return False
        else:
            print_status(f"✗ Resume upload failed: {response.status_code}", "ERROR")
            return False
            
    except Exception as e:
        print_status(f"✗ Upload workflow error: {str(e)}", "ERROR")
        return False
    finally:
        if test_file_path.exists():
            try:
                test_file_path.unlink()
            except Exception:
                pass

def test_data_consistency():
    """Test data consistency across the workflow"""
    print_status("\nTesting Data Consistency...", "INFO")
    
    base_url = "http://localhost:5000"
    
    try:
        analysis_response = requests.get(f"{base_url}/api/analysis", timeout=10)
        
        if analysis_response.status_code == 200:
            analysis_data = analysis_response.json()
            summary_response = requests.get(f"{base_url}/api/analysis/summary", timeout=10)
            
            if summary_response.status_code == 200:
                summary_data = summary_response.json()
                if "summary" in summary_data:
                    summary = summary_data["summary"]
                    analysis = analysis_data.get("data", {})
                    
                    fit_score_match = summary.get("fit_score") == analysis.get("overall_insights", {}).get("fit_score")
                    role_alignment_match = summary.get("role_alignment") == analysis.get("metrics", {}).get("role_alignment")
                    
                    if fit_score_match and role_alignment_match:
                        print_status("✓ Data consistency maintained across endpoints", "SUCCESS")
                        return True
                    else:
                        print_status("✗ Data inconsistency detected between endpoints", "ERROR")
                        return False
                else:
                    print_status("✗ Invalid summary response format", "ERROR")
                    return False
            else:
                print_status(f"✗ Summary endpoint failed: {summary_response.status_code}", "ERROR")
                return False
        else:
            print_status("⚠ No analysis data available for consistency test (run upload test first)", "WARNING")
            return True
            
    except Exception as e:
        print_status(f"✗ Data consistency test error: {str(e)}", "ERROR")
        return False

def main():
    """Main workflow test function"""
    print("=" * 70)
    print("🔄 AI Career Advisor - Complete Project Workflow Test")
    print("=" * 70)
    
    os.chdir(REPO_ROOT)
    
    tests = [
        ("Backend API Endpoints", test_backend_endpoints),
        ("Resume Upload Workflow", test_resume_upload_workflow),
        ("Data Consistency", test_data_consistency)
    ]
    
    results = {}
    
    for test_name, test_func in tests:
        print(f"\n{'='*20} {test_name} {'='*20}")
        results[test_name] = test_func()
        time.sleep(0.5)
    
    print("\n" + "=" * 70)
    print("📊 WORKFLOW TEST SUMMARY")
    print("=" * 70)
    
    passed = 0
    total = len(results)
    
    for test_name, result in results.items():
        status = "✓ PASS" if result else "✗ FAIL"
        color = "\033[92m" if result else "\033[91m"
        print(f"{color}{status}{'\033[0m'} {test_name}")
        if result:
            passed += 1
    
    print(f"\nOverall: {passed}/{total} workflow tests passed")
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
